import locale
from datetime import datetime
import pandas as pd
from docx import Document
from docx2pdf import convert
import smtplib
import pywhatkit
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import os
import re

# Dicionário de meses em português
MESES_PT = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}

# Configura o locale
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.utf8')
except:
    locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')

# Funções auxiliares
def limpar_nome(nome):
    nome = re.sub(r'[\\/:"*?<>|]', "", nome)
    return nome.strip()

def valor_por_extenso(valor):
    from num2words import num2words
    reais = int(valor)
    centavos = int(round((valor - reais) * 100))
    reais_extenso = num2words(reais, lang='pt-br')
    centavos_extenso = num2words(centavos, lang='pt-br')
    return f"{reais_extenso} reais e {centavos_extenso} centavos"

def substituir_texto(doc, substituicoes):
    for paragrafo in doc.paragraphs:
        for chave, valor in substituicoes.items():
            if chave in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(chave, valor)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, valor in substituicoes.items():
                    if chave in celula.text:
                        celula.text = celula.text.replace(chave, valor)

def calcular_vencimento():
    hoje = datetime.now()
    mes_vencimento = hoje.month + 1 if hoje.month < 12 else 1
    ano_vencimento = hoje.year if hoje.month < 12 else hoje.year + 1
    vencimento_data = datetime(ano_vencimento, mes_vencimento, 10)
    vencimento_formatado = f"10 de {MESES_PT[mes_vencimento]} de {ano_vencimento}"
    return vencimento_data, vencimento_formatado

def calcular_mes_referencia(referencia_avancada, vencimento_data):
    referencia_avancada = str(referencia_avancada).strip().lower()
    if referencia_avancada == "sim":
        return f"{MESES_PT[vencimento_data.month]}/{vencimento_data.year}", MESES_PT[vencimento_data.month]
    else:
        mes_anterior = vencimento_data.month - 1 if vencimento_data.month > 1 else 12
        ano_anterior = vencimento_data.year if vencimento_data.month > 1 else vencimento_data.year - 1
        return f"{MESES_PT[mes_anterior]}/{ano_anterior}", MESES_PT[mes_anterior]

def enviar_email(destinatario, nome_filiado, anexo_pdf):
    remetente = "abcohabs@uol.com.br"
    senha = "wamaneb2020"  # <--- Coloque sua senha

    servidor = "smtps.uol.com.br"
    porta = 587

    msg = MIMEMultipart()
    msg['From'] = remetente
    msg['To'] = destinatario
    msg['Subject'] = "Taxa de filiação à ABC"

    corpo_email = f"""\
Prezado(a) {nome_filiado},

Segue em anexo o ofício referente à taxa de filiação à ABC.

Por favor, qualquer dúvida estou à disposição.

Atenciosamente,
Associação Brasileira de Cohabs e Agentes Públicos de Habitação (ABC)
"""
    msg.attach(MIMEText(corpo_email, 'plain'))

    with open(anexo_pdf, 'rb') as f:
        pdf_part = MIMEApplication(f.read(), _subtype="pdf")
        pdf_part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(anexo_pdf))
        msg.attach(pdf_part)

    with smtplib.SMTP(servidor, porta) as smtp:
        smtp.starttls()
        smtp.login(remetente, senha)
        smtp.send_message(msg)

# --- Começa aqui ---
hoje = datetime.now()
ano_atual = hoje.year
mes_atual = MESES_PT[hoje.month]

# Lê planilha
filiados_df = pd.read_excel('filiados.xlsx')

# Cria pastas do mês atual
pasta_oficios = f"Oficios/{mes_atual}"
pasta_pdfs = f"PDFs/{mes_atual}"
os.makedirs(pasta_oficios, exist_ok=True)
os.makedirs(pasta_pdfs, exist_ok=True)

# Atualiza Número para nova planilha
filiados_df['Numero_Atualizado'] = filiados_df['Numero_Inicial'] + 1

# Lista para relatório
relatorio_envios = []

# Processa cada filiado
for idx, row in filiados_df.iterrows():
    filiado = row['Filiado']
    presidente = row['Presidente']
    valor_taxa = row['Valor_Taxa']
    referencia_avancada = row['Referencia_Avancada']
    email_destinatario = str(row['Email']).strip()

    numero_oficio_filiado = int(row['Numero_Inicial']) + 1
    numero_oficio = f"{numero_oficio_filiado:03}/{ano_atual}"

    valor_num = f"R$ {valor_taxa:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    valor_extenso = valor_por_extenso(valor_taxa)

    data_emissao = hoje.strftime('%d de %B de %Y').capitalize()

    vencimento_data, vencimento_formatado = calcular_vencimento()
    mes_referencia, mes_taxa = calcular_mes_referencia(referencia_avancada, vencimento_data)

    doc = Document('modelo_oficio.docx')
    substituicoes = {
        "{{DATA_EMISSAO}}": data_emissao,
        "{{NUMERO_OFICIO}}": numero_oficio,
        "{{FILIAL}}": filiado,
        "{{PRESIDENTE}}": presidente,
        "{{MES_REFERENCIA}}": mes_referencia,
        "{{VENCIMENTO}}": vencimento_formatado,
        "{{VALOR_NUM}}": valor_num,
        "{{VALOR_EXTENSO}}": valor_extenso,
        "{{MES_TAXA}}": mes_taxa
    }
    substituir_texto(doc, substituicoes)

    nome_base = limpar_nome(filiado)
    caminho_docx = os.path.join(pasta_oficios, f'Oficio_{nome_base}.docx')
    caminho_pdf = os.path.join(pasta_pdfs, f'Oficio_{nome_base}.pdf')

    doc.save(caminho_docx)
    convert(caminho_docx, caminho_pdf)

    status_envio = ""
    if pd.isna(email_destinatario) or email_destinatario == '' or email_destinatario.lower() == 'nan':
        status_envio = "⚠️ Não enviado (sem e-mail)"
        print(f"⚠️ {filiado} sem e-mail.")
    else:
        try:
            enviar_email(email_destinatario, filiado, caminho_pdf)
            status_envio = "✅ Enviado"
            print(f"✅ E-mail enviado para {email_destinatario}")
        except Exception as e:
            status_envio = f"❌ Erro: {str(e)}"
            print(f"❌ Erro ao enviar para {email_destinatario}: {e}")

    relatorio_envios.append({
        "Data_Hora": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "Filiado": filiado,
        "Email": email_destinatario,
        "Status": status_envio
    })

# Salva nova planilha
filiados_df[['Filiado', 'Presidente', 'Valor_Taxa', 'Numero_Atualizado', 'Referencia_Avancada', 'Email']].rename(
    columns={"Numero_Atualizado": "Numero_Inicial"}
).to_excel('filiados_atualizado.xlsx', index=False)

# Salva relatório
pd.DataFrame(relatorio_envios).to_csv('relatorio_envio.csv', index=False, sep=';')

# Conta quantos ofícios e quantos e-mails foram enviados com sucesso
total_oficios = len(filiados_df)
total_emails = sum(1 for r in relatorio_envios if r["Status"].startswith("✅"))

# Mensagem detalhada
mensagem_whats = (
    f"✅ Sistema de Ofícios ABC finalizado!\n"
    f"📄 {total_oficios} ofícios gerados\n"
    f"📧 {total_emails} e-mails enviados\n"
    f"📁 Relatório salvo em: {os.path.abspath('relatorio_envio.csv')}"
)

# Número no formato internacional
meu_numero = "+5519993457220"

# Envia instantaneamente (requer estar logado no WhatsApp Web)
pywhatkit.sendwhatmsg_instantly(meu_numero, mensagem_whats)

print(f"✅ Notificação enviada para WhatsApp: {meu_numero}")

print("\n✅ Oficios, PDFs, e relatórios do mês foram gerados com sucesso!")
